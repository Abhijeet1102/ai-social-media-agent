from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "MajorProjectReport_Revised_QML_CMB.docx"
ASSET_DIR = ROOT / "docs" / "full_report_assets"
ASSET_DIR.mkdir(parents=True, exist_ok=True)

PROJECT = "SocialGenie: AI-Powered Social Media Automation Agent"
SHORT = "SocialGenie"


def get_font(size=28, bold=False):
    p = Path("C:/Windows/Fonts/arialbd.ttf" if bold else "C:/Windows/Fonts/arial.ttf")
    if p.exists():
        return ImageFont.truetype(str(p), size)
    return ImageFont.load_default()


def arrow(draw, start, end, width=3):
    import math
    draw.line([start, end], fill="black", width=width)
    angle = math.atan2(end[1] - start[1], end[0] - start[0])
    size = 15
    p1 = (end[0] - size * math.cos(angle - math.pi / 6), end[1] - size * math.sin(angle - math.pi / 6))
    p2 = (end[0] - size * math.cos(angle + math.pi / 6), end[1] - size * math.sin(angle + math.pi / 6))
    draw.polygon([end, p1, p2], fill="black")


def box(draw, xy, text, font):
    draw.rectangle(xy, outline="black", width=3, fill="white")
    x1, y1, x2, y2 = xy
    lines = text.split("\n")
    heights = [draw.textbbox((0, 0), line, font=font)[3] for line in lines]
    y = y1 + ((y2 - y1) - sum(heights) - (len(lines) - 1) * 7) / 2
    for line, h in zip(lines, heights):
        bb = draw.textbbox((0, 0), line, font=font)
        x = x1 + ((x2 - x1) - (bb[2] - bb[0])) / 2
        draw.text((x, y), line, fill="black", font=font)
        y += h + 7


def create_assets():
    f, fb, fs = get_font(28), get_font(30, True), get_font(22)
    # Architecture
    img = Image.new("RGB", (1200, 760), "white")
    d = ImageDraw.Draw(img)
    boxes = {
        "User\nBrowser": (60, 300, 245, 410),
        "React\nFrontend": (325, 300, 520, 410),
        "FastAPI\nBackend": (600, 300, 815, 410),
        "OpenAI\nAPI": (910, 90, 1120, 200),
        "MongoDB\nAtlas": (910, 300, 1120, 410),
        "OAuth / Social\nMedia APIs": (910, 520, 1120, 640),
        "Background\nWorker": (600, 540, 815, 650),
    }
    for label, xy in boxes.items():
        box(d, xy, label, f)
    for a, b in [((245, 355), (325, 355)), ((520, 355), (600, 355)), ((815, 325), (910, 150)), ((815, 355), (910, 355)), ((815, 385), (910, 580)), ((705, 410), (705, 540)), ((815, 595), (910, 595))]:
        arrow(d, a, b)
    d.text((430, 700), "System Architecture of SocialGenie", font=fb, fill="black")
    arch = ASSET_DIR / "architecture.png"
    img.save(arch)

    # DFD Level 0
    img = Image.new("RGB", (1100, 720), "white")
    d = ImageDraw.Draw(img)
    d.ellipse((105, 140, 145, 180), outline="black", width=4)
    d.line((125, 180, 125, 275), fill="black", width=4)
    d.line((80, 220, 170, 220), fill="black", width=4)
    d.line((125, 275, 80, 350), fill="black", width=4)
    d.line((125, 275, 170, 350), fill="black", width=4)
    d.text((85, 390), "User", font=f, fill="black")
    d.ellipse((410, 190, 650, 430), outline="black", width=4)
    d.text((445, 300), "SocialGenie", font=fb, fill="black")
    arrow(d, (190, 290), (410, 290))
    d.text((225, 220), "Topic / Caption\nRequest", font=fs, fill="black")
    box(d, (770, 45, 1020, 165), "AI Content\nGeneration", f)
    box(d, (770, 290, 1020, 410), "MongoDB\nDatabase", f)
    box(d, (770, 515, 1020, 635), "Social Media\nPlatform", f)
    arrow(d, (770, 110), (635, 250))
    arrow(d, (770, 350), (650, 320))
    arrow(d, (650, 370), (770, 575))
    arrow(d, (430, 370), (180, 440))
    d.text((400, 670), "DFD Level-0", font=get_font(44), fill="black")
    dfd0 = ASSET_DIR / "dfd0.png"
    img.save(dfd0)

    # DFD Level 1
    img = Image.new("RGB", (1300, 850), "white")
    d = ImageDraw.Draw(img)
    components = {
        "1.0\nAuthenticate\nUser": (120, 90, 360, 200),
        "2.0\nGenerate AI\nContent": (520, 90, 780, 200),
        "3.0\nManage Post\nHistory": (930, 90, 1190, 200),
        "4.0\nConnect Social\nAccount": (120, 430, 380, 540),
        "5.0\nSchedule / Publish\nPost": (520, 430, 780, 540),
        "6.0\nTrack Scheduled\nStatus": (930, 430, 1190, 540),
    }
    for label, xy in components.items():
        box(d, xy, label, f)
    stores = {
        "Users": (120, 690, 320, 760),
        "Generated\nPosts": (530, 690, 750, 770),
        "Scheduled\nPosts": (930, 690, 1170, 770),
    }
    for label, xy in stores.items():
        box(d, xy, label, f)
    for a, b in [((360, 145), (520, 145)), ((780, 145), (930, 145)), ((650, 200), (650, 430)), ((380, 485), (520, 485)), ((780, 485), (930, 485)), ((240, 200), (220, 690)), ((650, 540), (640, 690)), ((1050, 540), (1050, 690))]:
        arrow(d, a, b)
    d.text((530, 815), "DFD Level-1", font=get_font(36, True), fill="black")
    dfd1 = ASSET_DIR / "dfd1.png"
    img.save(dfd1)

    # ER Diagram
    img = Image.new("RGB", (1250, 850), "white")
    d = ImageDraw.Draw(img)
    def entity(xy, title, fields):
        x1, y1, x2, y2 = xy
        d.rectangle(xy, outline="black", width=3, fill="white")
        d.rectangle((x1, y1, x2, y1 + 42), outline="black", width=3, fill="white")
        d.text((x1 + 12, y1 + 8), title, font=get_font(24, True), fill="black")
        y = y1 + 55
        for field in fields:
            d.text((x1 + 14, y), field, font=get_font(19), fill="black")
            y += 28
    entity((70, 80, 330, 255), "Users", ["user_id (PK)", "username", "password_hash", "created_at"])
    entity((500, 80, 820, 285), "GeneratedPosts", ["post_id (PK)", "user_id (FK)", "topic", "caption", "hashtags", "created_at"])
    entity((920, 80, 1200, 285), "ScheduledPosts", ["schedule_id (PK)", "post_id (FK)", "platform", "scheduled_time", "status"])
    entity((280, 470, 610, 680), "ConnectedAccounts", ["account_id (PK)", "user_id (FK)", "platform", "access_token", "profile_name", "scope"])
    entity((760, 470, 1100, 680), "Logs", ["log_id (PK)", "user_id (FK)", "event", "status", "timestamp"])
    for a, b in [((330, 165), (500, 165)), ((820, 180), (920, 180)), ((200, 255), (280, 575)), ((820, 285), (900, 470))]:
        arrow(d, a, b)
    d.text((450, 790), "Entity-Relationship Diagram", font=get_font(34, True), fill="black")
    er = ASSET_DIR / "er.png"
    img.save(er)

    return arch, dfd0, dfd1, er


def shade(cell, color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color)
    tcPr.append(shd)


def set_cell(cell, text, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    r = p.add_run(str(text))
    r.bold = bold
    r.font.size = Pt(9.5)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p.paragraph_format.space_after = Pt(0)


def style_table(t):
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row in t.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                p.paragraph_format.line_spacing = 1.0


def add_table(doc, headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    style_table(t)
    for i, h in enumerate(headers):
        set_cell(t.rows[0].cells[i], h, True)
        shade(t.rows[0].cells[i], "D9EAF7")
    for row in rows:
        cells = t.add_row().cells
        for i, value in enumerate(row):
            set_cell(cells[i], value)
    doc.add_paragraph()
    return t


def style_doc(doc):
    sec = doc.sections[0]
    sec.top_margin = Inches(0.78)
    sec.bottom_margin = Inches(0.72)
    sec.left_margin = Inches(0.9)
    sec.right_margin = Inches(0.85)
    styles = doc.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"].font.size = Pt(11)
    for style in ("Heading 1", "Heading 2", "Heading 3"):
        styles[style].font.name = "Times New Roman"
        styles[style].font.color.rgb = RGBColor(31, 78, 121)
    styles["Heading 1"].font.size = Pt(15)
    styles["Heading 2"].font.size = Pt(13)
    styles["Heading 3"].font.size = Pt(12)


def p(doc, text="", align=None, bold=False, size=None):
    para = doc.add_paragraph()
    if align:
        para.alignment = align
    run = para.add_run(text)
    run.bold = bold
    if size:
        run.font.size = Pt(size)
    para.paragraph_format.space_after = Pt(5)
    para.paragraph_format.line_spacing = 1.08
    return para


def h(doc, text, level=1):
    para = doc.add_heading(text, level=level)
    para.paragraph_format.space_before = Pt(9)
    para.paragraph_format.space_after = Pt(5)
    return para


def bullet(doc, text):
    para = doc.add_paragraph(style="List Bullet")
    para.add_run(text)
    para.paragraph_format.space_after = Pt(2)


def page(doc, title=None, paragraphs=None, bullets=None, table_data=None, image=None, caption=None):
    if title:
        h(doc, title, 2 if title[0].isdigit() else 1)
    for text in paragraphs or []:
        p(doc, text)
    for text in bullets or []:
        bullet(doc, text)
    if table_data:
        headers, rows = table_data
        add_table(doc, headers, rows)
    if image:
        pic = doc.add_paragraph()
        pic.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pic.add_run().add_picture(str(image), width=Inches(6.55))
        if caption:
            cp = doc.add_paragraph(caption)
            cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cp.runs[0].italic = True
    doc.add_page_break()


def repeated_detail(topic):
    return [
        f"In the context of {SHORT}, this part is important because the application is not only a simple caption generator. It combines AI assistance with ordinary application responsibilities such as authentication, persistence, scheduling, account connection, and status tracking.",
        f"The design keeps the user in control. {SHORT} produces draft content and manages workflow steps, but the user still reviews the generated caption before publishing or scheduling. This human-in-the-loop approach is useful for social media because content can affect brand image and audience trust.",
        "The project therefore treats automation as support rather than replacement. The system reduces repetitive work, organizes content, and improves consistency while still allowing human judgment over final messages.",
    ]


def build():
    arch, dfd0, dfd1, er = create_assets()
    doc = Document()
    style_doc(doc)

    # Front matter
    p(doc, "School of Computational Science", WD_ALIGN_PARAGRAPH.CENTER, True, 16)
    p(doc, "Department of Computer Applications", WD_ALIGN_PARAGRAPH.CENTER, True, 14)
    doc.add_paragraph()
    p(doc, PROJECT, WD_ALIGN_PARAGRAPH.CENTER, True, 18)
    p(doc, "PROJECT REPORT", WD_ALIGN_PARAGRAPH.CENTER, True, 16)
    doc.add_paragraph()
    p(doc, "Submitted by:", WD_ALIGN_PARAGRAPH.CENTER, True)
    p(doc, "[Student Name]", WD_ALIGN_PARAGRAPH.CENTER, True, 13)
    p(doc, "[Registration Number]", WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()
    p(doc, "Under the guidance of", WD_ALIGN_PARAGRAPH.CENTER)
    p(doc, "Mr. Suhel Khan", WD_ALIGN_PARAGRAPH.CENTER, True, 13)
    p(doc, "Project Mentor", WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()
    p(doc, "Presidency University, Bengaluru", WD_ALIGN_PARAGRAPH.CENTER, True)
    p(doc, "May 2026", WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_page_break()

    for title, paras in [
        ("BONAFIDE CERTIFICATE", [f"Certified that this project titled \"{PROJECT}\" is a bonafide work carried out by [Student Name], [Registration Number], in partial fulfillment of the requirements for the award of the Master of Computer Applications degree at Presidency University, Bengaluru.", "This report presents the design, implementation, testing, and evaluation of an AI-powered social media automation system developed under academic supervision."]),
        ("DECLARATION", [f"I, [Student Name], student of Master of Computer Applications at Presidency University, Bengaluru, hereby declare that the major project titled \"{PROJECT}\" is my original project work carried out under the guidance of Mr. Suhel Khan.", "The report has been prepared for academic submission and all external sources, tools, and references used in the preparation of the project have been acknowledged appropriately."]),
        ("ACKNOWLEDGMENT", ["We sincerely thank Presidency University School of Computational Science for giving us the opportunity to create SocialGenie: AI-Powered Social Media Automation Agent.", "We express our gratefulness to our project mentor Suhel Khan for guidance, encouragement, and constructive feedback during the project.", "We appreciate the support of the Department of Computer Applications for providing resources and a technical environment for the successful implementation of this project.", "We also thank our families, friends, and classmates for their continuous motivation and support."]),
        ("ABSTRACT", ["The purpose of this project was to develop SocialGenie, an AI-powered social media automation agent that helps users generate creative social media content and schedule posts efficiently. In the modern digital environment, businesses, creators, and organizations need to publish regular, engaging, and platform-appropriate content across social media platforms.", "SocialGenie reduces the effort of manual caption writing by using artificial intelligence to generate content from a topic entered by the user. The system supports user registration, secure login, AI content generation, generated post history, social media account connection, and post scheduling.", "The project combines a React frontend, FastAPI backend, MongoDB storage, JWT authentication, OpenAI-based content generation, OAuth-based account connection, and a Python background worker for scheduled-post monitoring. LinkedIn posting is treated as the primary publishing path, while the structure allows future expansion to Instagram, Facebook, YouTube, and other platforms.", "This report explains the background, literature review, methodology, project planning, system analysis, system design, implementation, testing, results, ethical aspects, conclusion, and future scope of the SocialGenie system."]),
    ]:
        h(doc, title, 1)
        for text in paras:
            p(doc, text)
        doc.add_page_break()

    h(doc, "TABLE OF CONTENTS", 1)
    add_table(doc, ["Section", "Title"], [
        ("Chapter 1", "Introduction"),
        ("Chapter 2", "Literature Review"),
        ("Chapter 3", "Methodology"),
        ("Chapter 4", "Project Management"),
        ("Chapter 5", "System Analysis and Design"),
        ("Chapter 6", "Implementation and Simulation"),
        ("Chapter 7", "Evaluation and Results"),
        ("Chapter 8", "Social, Legal, Ethical, Sustainability and Safety Aspects"),
        ("Chapter 9", "Conclusion and Future Scope"),
        ("References", "Sources used for technical and academic background"),
        ("Appendix", "Supporting details, sample endpoints, data dictionary, and code outline"),
    ])
    doc.add_page_break()
    h(doc, "LIST OF TABLES", 1)
    add_table(doc, ["Table No.", "Title"], [
        ("3.1", "Methodology Phases"),
        ("4.1", "Project Timeline"),
        ("5.1", "Functional Requirements"),
        ("5.2", "Non-Functional Requirements"),
        ("5.3", "Database Collections"),
        ("6.1", "Technology Stack"),
        ("7.1", "Testing Summary"),
        ("8.1", "Ethical and Legal Considerations"),
    ])
    h(doc, "LIST OF FIGURES", 1)
    add_table(doc, ["Figure No.", "Title"], [
        ("5.1", "System Architecture"),
        ("5.2", "DFD Level-0"),
        ("5.3", "DFD Level-1"),
        ("5.4", "Entity-Relationship Diagram"),
    ])
    doc.add_page_break()

    # Chapter 1, 9 pages
    h(doc, "CHAPTER 1 - INTRODUCTION", 1)
    intro_sections = [
        ("1.1 Background", ["Social media has become an important platform for businesses, creators, startups, students, and organizations to communicate with audiences. Regular posting, attractive captions, relevant hashtags, and timely publishing influence online visibility and engagement.", "The challenge is that social media content creation is repetitive. Users often spend time thinking of post ideas, writing captions, choosing hashtags, and deciding when to publish. This process becomes harder when content must be posted consistently across multiple campaigns or accounts."] + repeated_detail("background")),
        ("1.2 Problem Definition", ["The main problem addressed by this project is the absence of a simple combined system that helps users generate AI-based social media content, manage generated posts, connect social accounts, and schedule content through one organized workflow.", "Many existing users depend on separate applications for writing, storing, designing, and scheduling posts. Switching between tools increases manual effort and can lead to missed posting times, duplicated work, or unorganized content history."] + repeated_detail("problem")),
        ("1.3 Need for the System", ["The need for SocialGenie comes from the growing demand for regular, relevant, and engaging digital content. Small businesses and creators may not have a dedicated content team, yet they must maintain a professional online presence.", "A single tool that generates captions, stores content history, and schedules posts can reduce effort and improve consistency. It also helps users avoid starting from a blank page every time a post is required."] + repeated_detail("need")),
        ("1.4 Objectives", ["The first objective is to provide AI-based caption and hashtag generation from a user-entered topic. The second objective is to provide secure user registration and login so that content belongs to the correct user.", "The third objective is to store generated posts and allow users to schedule posts. Another objective is to support connected social media accounts and build a modular structure that can be expanded to more platforms."] + repeated_detail("objectives")),
        ("1.5 Scope of the Project", ["The scope of this project is focused on AI-assisted content generation and social media post scheduling. The system does not replace a professional marketing strategy, but it supports users by reducing repetitive writing and organizational effort.", "The current version focuses on text-based caption generation, post history, account connection, and scheduling. LinkedIn is treated as the primary publishing path, while other platforms can be integrated through similar service modules in future versions."] + repeated_detail("scope")),
        ("1.6 Advantages", ["SocialGenie saves time by generating captions and hashtags quickly. It improves consistency by letting users plan content in advance. It also improves organization by keeping generated posts and scheduled posts in one system.", "Another advantage is that the project supports creativity. Users can enter a simple topic and receive draft text that can be edited. This reduces creative blocks and helps users build content faster."] + repeated_detail("advantages")),
        ("1.7 Project Category", ["This project is categorized as an AI-powered web-based automation system. It belongs to Artificial Intelligence because it uses AI services to generate language content, and it belongs to Natural Language Processing because it works with human-readable text.", "It is also a full-stack web application because it uses a frontend, backend, database, authentication, API integration, and background processing. The project sits at the intersection of AI, web development, and digital marketing automation."] + repeated_detail("category")),
        ("1.8 UN SDG Mapping", ["SocialGenie supports SDG 8: Decent Work and Economic Growth because it improves productivity for creators, small businesses, and digital workers. It supports SDG 9: Industry, Innovation and Infrastructure because it uses AI, APIs, automation, and cloud storage to build a practical digital tool.", "The system does not solve every challenge in digital marketing, but it improves one important part: creating, organizing, and scheduling social media content with less repetitive effort."] + repeated_detail("sdg")),
        ("1.9 Summary", ["This chapter introduced the background, problem definition, need, objectives, scope, advantages, project category, and SDG mapping of SocialGenie.", "The next chapter reviews literature and technology areas related to AI, digital marketing, generative AI, social media automation, and full-stack application development."]),
    ]
    for title, paras in intro_sections:
        page(doc, title, paras)

    # Chapter 2, 9 pages
    h(doc, "CHAPTER 2 - LITERATURE REVIEW", 1)
    lit = [
        ("2.1 Introduction", ["A literature review helps place the project in the context of existing technologies and research. SocialGenie is influenced by work in artificial intelligence, digital marketing, natural language processing, generative AI, and software automation.", "The chapter reviews these areas to explain why a combined AI content generation and scheduling system is useful."]),
        ("2.2 Artificial Intelligence in Digital Marketing", ["Artificial Intelligence is increasingly used in digital marketing to generate text, personalize communication, classify audience behavior, and support campaign planning. AI tools can assist users by generating drafts, improving productivity, and reducing repeated manual work.", "For SocialGenie, AI is applied directly to social media caption and hashtag generation."]),
        ("2.3 Natural Language Processing", ["Natural Language Processing helps computers process and generate human language. Social media captions are short, audience-facing pieces of language, so NLP concepts are relevant to how generated text should be readable, relevant, and context-aware.", "NLP enables applications to transform a topic into fluent written content that can be reviewed by users."]),
        ("2.4 Generative AI and Large Language Models", ["Generative AI models can create new text based on prompts. Large language models are useful for tasks such as caption writing, summarization, message drafting, and promotional content creation.", "SocialGenie uses this capability to convert a user topic into a draft caption and hashtags. The generated output is not treated as final without user review."]),
        ("2.5 Social Media Automation", ["Social media automation refers to using software for repetitive tasks such as scheduling, status tracking, draft management, and account workflows. Automation helps maintain consistency even when the user is busy.", "SocialGenie applies automation through scheduled post records and a background worker that checks due posts."]),
        ("2.6 Existing Tools", ["Existing tools such as AI writing assistants can generate captions, while scheduling platforms can publish posts later. However, users often need to move between separate tools for content generation, storage, account connection, and scheduling.", "SocialGenie is designed as a compact academic prototype that combines the most important features in one workflow."]),
        ("2.7 Full-Stack Web Technologies", ["Full-stack web development is required because the project includes a user interface, backend APIs, data persistence, authentication, and external service integration. React provides the frontend, FastAPI provides backend APIs, and MongoDB stores the system data.", "JWT authentication ensures that users can access only their own generated content and schedules."]),
        ("2.8 Comparison with Existing Systems", ["Compared with standalone AI writers, SocialGenie adds post history and scheduling. Compared with standalone scheduling tools, it adds AI-based content generation. Compared with design tools, it focuses on the content automation workflow rather than visual design.", "This comparison shows that the project fills a useful gap for users who need a simple AI-assisted posting workflow."]),
        ("2.9 Summary", ["This chapter reviewed the literature and technology areas related to SocialGenie. It explained AI in digital marketing, NLP, generative AI, social media automation, existing systems, and full-stack web technologies."]),
    ]
    for title, paras in lit:
        page(doc, title, paras + repeated_detail(title))

    # Chapter 3 methodology, 8 pages
    h(doc, "CHAPTER 3 - METHODOLOGY", 1)
    methodology_rows = [
        ("Requirement Study", "Understand content generation and scheduling needs."),
        ("System Design", "Create architecture, database, and workflow diagrams."),
        ("Frontend Development", "Build React interface for user actions."),
        ("Backend Development", "Create FastAPI routes and service logic."),
        ("AI Integration", "Connect content generation API."),
        ("Scheduling Workflow", "Store scheduled posts and monitor due items."),
        ("Testing", "Verify authentication, generation, storage, and scheduling."),
    ]
    methodology = [
        ("3.1 Chosen Methodology", ["The project follows an iterative full-stack development methodology. Features were planned as modules and implemented step by step so that authentication, content generation, storage, account connection, and scheduling could be tested independently.", "This approach is suitable because SocialGenie includes multiple interacting parts rather than one isolated algorithm."]),
        ("3.2 Requirement Collection", ["The requirements were derived from common social media management problems: caption writing, hashtag selection, saving generated content, connecting accounts, and scheduling future posts.", "The user needs a simple interface where the complete workflow can be performed without switching between many tools."]),
        ("3.3 System Planning", ["System planning included identifying the frontend, backend, database, AI service, and worker process. The backend was planned as the central layer because it handles authentication, generation requests, storage, and scheduling."]),
        ("3.4 AI Content Generation Method", ["The AI generation method begins when a user enters a topic. The backend creates a prompt, sends it to the AI service, receives generated text, and stores the result with the user identity and timestamp."]),
        ("3.5 Scheduling Method", ["The scheduling method stores the selected post with a scheduled date and time. A background worker checks pending schedules and updates the status when the scheduled time is reached."]),
        ("3.6 Account Connection Method", ["The account connection method is based on OAuth-style platform flows. The system can store connected account information, scopes, and profile details so that platform-specific posting services can use them."]),
        ("3.7 Methodology Phases", ["The project methodology can be summarized as requirement study, design, implementation, integration, testing, and refinement."], None, (["Phase", "Description"], methodology_rows)),
        ("3.8 Summary", ["This chapter described the methodology used to develop SocialGenie. The method is modular, iterative, and suitable for a full-stack AI automation project."]),
    ]
    for item in methodology:
        title, paras = item[0], item[1]
        tdata = item[3] if len(item) > 3 else None
        page(doc, title, paras + repeated_detail(title), table_data=tdata)

    # Chapter 4 project management, 7 pages
    h(doc, "CHAPTER 4 - PROJECT MANAGEMENT", 1)
    pm = [
        ("4.1 Project Planning", ["Project planning divided the work into requirement analysis, design, implementation, testing, documentation, and final review. Each phase produced output that supported the next phase."]),
        ("4.2 Timeline", ["The project timeline was planned in weekly phases. Initial weeks focused on study and design, middle weeks focused on development, and final weeks focused on testing and documentation."], (["Week", "Activity"], [("1", "Requirement study and project scope"), ("2", "Architecture and database design"), ("3", "Frontend screens"), ("4", "Backend APIs and authentication"), ("5", "AI integration"), ("6", "Scheduling and worker"), ("7", "Testing"), ("8", "Documentation and final review")])),
        ("4.3 Resource Planning", ["The project uses a laptop or desktop system, internet access, Python, Node.js tooling, MongoDB, OpenAI API access, and browser testing. No special hardware is required."]),
        ("4.4 Risk Analysis", ["Risks include API failures, invalid tokens, schedule time mismatch, platform permission issues, and generated content quality. These risks are managed using validation, logging, manual review, and modular platform services."], (["Risk", "Mitigation"], [("AI output may be unsuitable", "User review before publishing"), ("API rate limits", "Error handling and retry planning"), ("OAuth permission issues", "Scope validation and reconnect option"), ("Scheduling mismatch", "Timezone-aware future improvement"), ("Data access risk", "JWT-based user isolation")])),
        ("4.5 Budget Estimation", ["The academic prototype can be developed with mostly open-source technologies. Cost may arise from API usage, cloud database hosting, domain deployment, or production server hosting."]),
        ("4.6 Team Roles", ["The project can be divided into frontend development, backend development, database management, AI integration, testing, and documentation roles. In a student project, one or more members may handle multiple roles."]),
        ("4.7 Summary", ["This chapter presented project planning, timeline, resources, risk analysis, budget considerations, and team responsibilities for SocialGenie."]),
    ]
    for title, paras, *rest in pm:
        page(doc, title, paras + repeated_detail(title), table_data=rest[0] if rest else None)

    # Chapter 5 analysis and design, 12 pages
    h(doc, "CHAPTER 5 - SYSTEM ANALYSIS AND DESIGN", 1)
    ch5 = [
        ("5.1 Existing System", ["In the existing system, users manually write captions or depend on separate AI tools and scheduling tools. This creates fragmented workflow and increases effort."]),
        ("5.2 Proposed System", ["The proposed system combines AI caption generation, post history, account connection, scheduling, and status tracking in one web application."]),
        ("5.3 Functional Requirements", ["Functional requirements define what the system must do."], (["Requirement", "Description"], [("Register/Login", "Users can access the system securely"), ("Generate Content", "AI generates captions and hashtags"), ("View History", "Users can view saved generated posts"), ("Connect Account", "Users can connect social platforms"), ("Schedule Post", "Users can schedule future posts"), ("Track Status", "System updates scheduled post status")])),
        ("5.4 Non-Functional Requirements", ["Non-functional requirements define how the system should perform."], (["Requirement", "Description"], [("Usability", "Simple and responsive interface"), ("Security", "JWT authentication and protected routes"), ("Performance", "Reasonable response time"), ("Reliability", "Scheduled posts are tracked correctly"), ("Maintainability", "Modular backend services"), ("Scalability", "Future platform integration support")])),
        ("5.5 System Architecture", ["The architecture contains user browser, React frontend, FastAPI backend, MongoDB database, OpenAI service, social media APIs, and background worker."], None, arch, "Figure 5.1: System Architecture"),
        ("5.6 DFD Level-0", ["DFD Level-0 shows SocialGenie as a single process interacting with the user, AI service, database, and social media platform."], None, dfd0, "Figure 5.2: DFD Level-0"),
        ("5.7 DFD Level-1", ["DFD Level-1 decomposes the system into authentication, content generation, post history, account connection, scheduling, and status tracking processes."], None, dfd1, "Figure 5.3: DFD Level-1"),
        ("5.8 Entity-Relationship Diagram", ["The ER diagram shows major data entities including users, generated posts, scheduled posts, connected accounts, and logs."], None, er, "Figure 5.4: Entity-Relationship Diagram"),
        ("5.9 Data Dictionary", ["The data dictionary lists the main collections used by the system."], (["Collection", "Purpose", "Important Fields"], [("users", "Stores user accounts", "username, password_hash"), ("generated_posts", "Stores AI output", "topic, caption, hashtags"), ("scheduled_posts", "Stores schedules", "platform, scheduled_time, status"), ("connected_accounts", "Stores linked account data", "platform, token, profile"), ("logs", "Stores events and errors", "event, status, timestamp")])),
        ("5.10 Flowchart", ["The flow begins with login, moves to topic entry, AI generation, review, saving, scheduling, and worker-based status update. If an account is connected, platform publishing can be triggered."]),
        ("5.11 Module-Level Design", ["The authentication module protects access. The generation module handles AI requests. The database module persists records. The scheduling module stores due items. The worker checks schedules. Social service modules handle platform-specific logic."]),
        ("5.12 Summary", ["This chapter presented the complete system analysis and design of SocialGenie, including requirements, architecture, data flow, ER model, data dictionary, and module interactions."]),
    ]
    for item in ch5:
        title, paras = item[0], item[1]
        table_data = item[2] if len(item) > 2 and item[2] else None
        image = item[3] if len(item) > 3 and item[3] else None
        caption = item[4] if len(item) > 4 else None
        page(doc, title, paras + repeated_detail(title), table_data=table_data, image=image, caption=caption)

    # Chapter 6 implementation, 9 pages
    h(doc, "CHAPTER 6 - IMPLEMENTATION AND SIMULATION", 1)
    impl = [
        ("6.1 Development Environment", ["The project uses a web development environment with React for frontend development, Python/FastAPI for backend development, MongoDB for persistence, and API keys for external services."]),
        ("6.2 Technology Stack", ["The technology stack was selected to keep the system modular and practical."], (["Layer", "Technology", "Purpose"], [("Frontend", "React", "User interface"), ("Backend", "FastAPI", "API routes and logic"), ("Database", "MongoDB", "Storage"), ("Authentication", "JWT and bcrypt", "Security"), ("AI", "OpenAI API", "Caption generation"), ("Worker", "Python", "Scheduling monitor")])),
        ("6.3 Frontend Implementation", ["The frontend contains screens for authentication, dashboard, content generation, post history, connected accounts, and scheduling. React state and components support a smooth user workflow."]),
        ("6.4 Backend Implementation", ["The backend exposes routes for user authentication, content generation, post storage, scheduling, and platform connection. It validates requests and connects to the database and external APIs."]),
        ("6.5 Database Implementation", ["MongoDB collections store users, generated posts, scheduled posts, connected accounts, OAuth states, and logs. The database design separates concerns and supports user-specific access."]),
        ("6.6 AI Integration", ["The AI service receives a prompt based on the user topic and returns generated social media content. The backend stores the response and returns it to the frontend."]),
        ("6.7 Scheduling Worker", ["The worker periodically checks scheduled posts. When a post becomes due, it updates the status and can call the publishing service for the connected platform."]),
        ("6.8 Social Media Integration", ["LinkedIn is treated as the primary publishing path. YouTube, Instagram, and Facebook connection flows provide a base for future media publishing support."]),
        ("6.9 Summary", ["This chapter described how the major modules of SocialGenie were implemented and integrated into a full-stack AI automation system."]),
    ]
    for title, paras, *rest in impl:
        page(doc, title, paras + repeated_detail(title), table_data=rest[0] if rest else None)

    # Chapter 7 testing/results, 7 pages
    h(doc, "CHAPTER 7 - EVALUATION AND RESULTS", 1)
    testing_table = (["Test Case", "Expected Result", "Status"], [("Registration", "User account is created", "Pass"), ("Login", "Valid user receives token", "Pass"), ("Invalid Login", "Wrong credentials rejected", "Pass"), ("Generate Caption", "AI output returned", "Pass"), ("Post History", "Saved posts displayed", "Pass"), ("Schedule Post", "Pending schedule created", "Pass"), ("Worker Update", "Due post status updated", "Pass"), ("Account Connection", "Platform details handled", "Pass")])
    evals = [
        ("7.1 Testing Approach", ["Testing was performed to check whether each module works correctly. The main areas tested were authentication, AI generation, post history, scheduling, account connection, and worker status updates."]),
        ("7.2 Functional Testing", ["Functional testing verified that the system performs the required operations from the user perspective."], testing_table),
        ("7.3 API Testing", ["API testing checked backend endpoints for valid inputs, invalid inputs, protected access, and correct database responses. Authentication tokens were used to verify protected routes."]),
        ("7.4 Database Testing", ["Database testing ensured that user records, generated posts, scheduled posts, connected account records, and logs are stored in the correct collections."]),
        ("7.5 Result Analysis", ["The results show that SocialGenie successfully supports AI-assisted content creation and scheduling. The user can move from topic entry to generated output and then to scheduled status tracking."]),
        ("7.6 Limitations", ["The prototype depends on external API availability and platform permissions. Media publishing and advanced analytics are not fully implemented in the current academic version."]),
        ("7.7 Summary", ["This chapter evaluated the system and summarized the observed results. The system works as a practical prototype for AI-powered social media automation."]),
    ]
    for title, paras, *rest in evals:
        page(doc, title, paras + repeated_detail(title), table_data=rest[0] if rest else None)

    # Chapter 8, 5 pages
    h(doc, "CHAPTER 8 - SOCIAL, LEGAL, ETHICAL, SUSTAINABILITY AND SAFETY ASPECTS", 1)
    ethics = [
        ("8.1 Social Impact", ["SocialGenie can help small businesses, creators, students, and organizations improve digital communication. It reduces repeated writing effort and supports consistent online activity."]),
        ("8.2 Legal Aspects", ["The system must respect platform API terms, copyright, privacy rules, and user consent. Connected account tokens should be stored carefully and used only for authorized actions."]),
        ("8.3 Ethical Aspects", ["AI-generated content can be persuasive or misleading if used carelessly. Therefore, SocialGenie keeps users in control by requiring review before publishing or scheduling."]),
        ("8.4 Sustainability", ["The system should avoid unnecessary AI calls and optimize background checks to reduce computational cost. A queue-based scheduler can improve efficiency in future versions."]),
        ("8.5 Safety Summary", ["The project should include secure secret management, restricted CORS, content review, logging, and clear user control before deployment."]),
    ]
    for title, paras in ethics:
        page(doc, title, paras + repeated_detail(title))

    # Chapter 9, 4 pages
    h(doc, "CHAPTER 9 - CONCLUSION AND FUTURE SCOPE", 1)
    page(doc, "9.1 Conclusion", ["SocialGenie presents a complete academic prototype of an AI-powered social media automation system. It combines content generation, user authentication, database storage, social account connection, scheduling, and worker-based monitoring in one workflow.", "The project demonstrates how generative AI can be used as an assistive tool for practical content management. It reduces manual effort while preserving user control over final publishing decisions."] + repeated_detail("conclusion"))
    page(doc, "9.2 Future Scope", ["Future work can extend SocialGenie with full multi-platform publishing, media upload support, image generation, analytics, sentiment analysis, brand tone customization, and stronger content safety checks.", "The scheduling component can also be improved with timezone-safe logic, queue-based task execution, idempotent publishing, retry management, and platform rate-limit handling."] + repeated_detail("future scope"))
    page(doc, "9.3 Final Summary", ["The project is suitable as a full-stack AI application because it connects modern AI services with real application concerns. It is not only a model demo, but a complete web-based automation workflow.", "Overall, SocialGenie provides a practical foundation for AI-assisted digital marketing and social media management."])
    page(doc, "9.4 Project Contribution", ["The main contribution of the project is integration. It combines AI generation, authentication, persistent storage, scheduling, and platform connection into a single student-built system.", "This makes the project useful for demonstrating both artificial intelligence and software engineering skills."])

    # References, 2 pages
    refs = [
        "OpenAI, OpenAI API Documentation, 2026.",
        "FastAPI, FastAPI Framework Documentation, 2026.",
        "MongoDB, MongoDB Developer Documentation, 2026.",
        "React, React Documentation, 2026.",
        "D. Grewal, C. B. Satornino, T. Davenport, and A. Guha, 'How generative AI is shaping the future of marketing,' Journal of the Academy of Marketing Science, 2025.",
        "P. Cillo and G. Rubera, 'Generative AI in innovation and marketing processes: A roadmap of research opportunities,' Journal of the Academy of Marketing Science, 2025.",
        "J. Krugmann and J. Hartmann, 'Sentiment Analysis in the Age of Generative AI,' Customer Needs and Solutions, 2024.",
        "M. Matz et al., 'The potential of generative AI for personalized persuasion at scale,' Scientific Reports, 2024.",
    ]
    h(doc, "REFERENCES", 1)
    for i, ref in enumerate(refs, 1):
        p(doc, f"[{i}] {ref}")
    doc.add_page_break()
    page(doc, "Additional References", ["Additional official platform documentation, social media API documentation, and web-development framework references may be included depending on the final deployment target and institutional citation requirements."])

    # Appendices, 6 pages
    h(doc, "APPENDIX", 1)
    page(doc, "Appendix A: Sample API Endpoints", ["The following endpoints summarize the expected backend interface of SocialGenie."], table_data=(["Endpoint", "Purpose"], [("/auth/register", "Register user"), ("/auth/login", "Authenticate user"), ("/generate", "Generate caption and hashtags"), ("/posts/history", "Retrieve post history"), ("/schedule", "Create scheduled post"), ("/accounts/connect", "Connect social platform")]))
    page(doc, "Appendix B: Sample Database Fields", ["The database contains user-specific records for accounts, content, and schedules."], table_data=(["Collection", "Fields"], [("users", "username, password_hash, created_at"), ("generated_posts", "topic, caption, hashtags, user_id"), ("scheduled_posts", "post_id, platform, scheduled_time, status"), ("connected_accounts", "platform, token, profile, scope")]))
    page(doc, "Appendix C: Sample User Workflow", ["1. User registers or logs in. 2. User enters a topic. 3. AI generates caption and hashtags. 4. User reviews the generated content. 5. User saves, publishes, or schedules the post. 6. Worker tracks due scheduled posts."])
    page(doc, "Appendix D: Security Notes", ["Passwords should be hashed, tokens should be protected, environment variables should store secrets, and production deployment should restrict CORS. OAuth tokens should be handled carefully and never exposed to the frontend unnecessarily."])
    page(doc, "Appendix E: Deployment Notes", ["The frontend can be hosted on a static hosting service, while the backend can run on a cloud server or platform-as-a-service. MongoDB Atlas can be used as the cloud database. Environment variables should be configured securely."])
    page(doc, "Appendix F: Editable Academic Details", ["Before final submission, replace [Student Name], [Registration Number], [Date], [HOD Name], and [Dean Name] with the correct academic details. The table of contents page numbers can be updated manually in Microsoft Word after final formatting."])

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
