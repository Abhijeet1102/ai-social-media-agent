import re
import zipfile
import docx

p = r"docs\MajorProjectReport_Revised_QML_CMB.docx"
d = docx.Document(p)
heads = [x.text for x in d.paragraphs if x.style.name.startswith("Heading") and x.text.strip()]
xml = zipfile.ZipFile(p).read("word/document.xml").decode("utf-8")
print("paragraphs", len(d.paragraphs))
print("tables", len(d.tables))
print("images", len(d.inline_shapes))
print("headings", len(heads))
print("first_heading", heads[0])
print("last_heading", heads[-1])
print("page_breaks", xml.count('type="page"'))
print("contains_old_cmb", bool(re.search("Cosmic Microwave|Quantum|Ramanujan", xml, re.I)))
