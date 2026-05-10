from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION_START
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "SocialGenie_Project_Report.docx"
ASSET_DIR = ROOT / "docs" / "report_assets"
ASSET_DIR.mkdir(parents=True, exist_ok=True)


PROJECT = "SocialGenie: AI-Powered Social Media Automation Agent"
SHORT = "SocialGenie"


def font(size=28, bold=False):
    candidates = [
        "C:/Windows/Fonts/arialbd.ttf" if bold else "C:/Windows/Fonts/arial.ttf",
        "C:/Windows/Fonts/calibrib.ttf" if bold else "C:/Windows/Fonts/calibri.ttf",
    ]
    for p in candidates:
        if Path(p).exists():
            return ImageFont.truetype(p, size)
    return ImageFont.load_default()


def draw_box(draw, xy, text, fnt, width=3):
    draw.rectangle(xy, outline="black", width=width, fill="white")
    x1, y1, x2, y2 = xy
    lines = text.split("\n")
    total_h = sum(draw.textbbox((0, 0), line, font=fnt)[3] for line in lines) + (len(lines)-1) * 8
    y = y1 + ((y2 - y1) - total_h) / 2
    for line in lines:
        bbox = draw.textbbox((0, 0), line, font=fnt)
        x = x1 + ((x2 - x1) - (bbox[2] - bbox[0])) / 2
        draw.text((x, y), line, fill="black", font=fnt)
        y += (bbox[3] - bbox[1]) + 8


def arrow(draw, start, end, width=3):
    draw.line([start, end], fill="black", width=width)
    import math
    angle = math.atan2(end[1]-start[1], end[0]-start[0])
    size = 16
    p1 = (end[0] - size*math.cos(angle - math.pi/6), end[1] - size*math.sin(angle - math.pi/6))
    p2 = (end[0] - size*math.cos(angle + math.pi/6), end[1] - size*math.sin(angle + math.pi/6))
    draw.polygon([end, p1, p2], fill="black")


def create_diagrams():
    # DFD Level 0
    img = Image.new("RGB", (1100, 720), "white")
    d = ImageDraw.Draw(img)
    f_big, f, f_small = font(34), font(30), font(24)
    # Actor
    d.ellipse((100, 135, 145, 180), outline="black", width=4)
    d.line((122, 180, 122, 280), fill="black", width=4)
    d.line((75, 220, 170, 220), fill="black", width=4)
    d.line((122, 280, 75, 365), fill="black", width=4)
    d.line((122, 280, 170, 365), fill="black", width=4)
    d.text((78, 395), "User", fill="black", font=f_big)
    d.ellipse((410, 185, 650, 425), outline="black", width=4, fill="white")
    d.text((445, 285), "SocialGenie", fill="black", font=f_big)
    arrow(d, (190, 285), (410, 285))
    d.text((225, 215), "Enter Topic /\nRequest Caption", fill="black", font=f_small)
    draw_box(d, (770, 45, 1025, 165), "AI Content\nGeneration", f)
    arrow(d, (770, 115), (640, 245))
    d.text((550, 80), "Generate\nCaption", fill="black", font=f_small)
    draw_box(d, (770, 290, 1025, 410), "MongoDB\nDatabase", f)
    arrow(d, (770, 350), (650, 315))
    d.text((555, 430), "Store / Retrieve\nPosts", fill="black", font=f_small)
    draw_box(d, (770, 515, 1025, 635), "Social Media\nPlatform", f)
    arrow(d, (650, 365), (770, 575))
    d.text((555, 550), "Schedule /\nPublish", fill="black", font=f_small)
    arrow(d, (430, 365), (180, 430))
    d.text((230, 470), "Generated Content /\nStatus", fill="black", font=f_small)
    d.text((410, 670), "DFD Level-0", fill="black", font=font(44))
    p0 = ASSET_DIR / "dfd_level_0.png"
    img.save(p0)

    # Architecture
    img = Image.new("RGB", (1200, 760), "white")
    d = ImageDraw.Draw(img)
    f = font(28)
    boxes = {
        "React\nFrontend": (80, 290, 280, 410),
        "FastAPI\nBackend": (430, 290, 650, 410),
        "OpenAI\nService": (800, 95, 1030, 210),
        "MongoDB\nDatabase": (800, 315, 1030, 430),
        "Background\nWorker": (430, 540, 650, 650),
        "Social Media\nAPIs": (800, 540, 1030, 650),
    }
    for text, xy in boxes.items():
        draw_box(d, xy, text, f)
    arrow(d, (280, 350), (430, 350))
    arrow(d, (650, 315), (800, 170))
    arrow(d, (650, 350), (800, 370))
    arrow(d, (540, 410), (540, 540))
    arrow(d, (650, 595), (800, 595))
    arrow(d, (430, 370), (280, 370))
    d.text((335, 315), "API\nRequests", fill="black", font=font(22))
    d.text((875, 705), "System Architecture", fill="black", font=font(30, True))
    p1 = ASSET_DIR / "system_architecture.png"
    img.save(p1)

    # ER diagram
    img = Image.new("RGB", (1200, 760), "white")
    d = ImageDraw.Draw(img)
    ent = font(26, True)
    attrs = font(20)
    def entity(xy, name, fields):
        draw_box(d, xy, name + "\n" + "\n".join(fields), attrs)
        x1, y1, x2, y2 = xy
        d.rectangle((x1, y1, x2, y1+38), outline="black", width=3, fill="white")
        tw = d.textbbox((0, 0), name, font=ent)
        d.text((x1+(x2-x1-(tw[2]-tw[0]))/2, y1+6), name, fill="black", font=ent)
    entity((80, 90, 330, 260), "Users", ["user_id (PK)", "username", "password_hash"])
    entity((470, 90, 760, 280), "GeneratedPosts", ["post_id (PK)", "user_id (FK)", "topic", "caption", "hashtags"])
    entity((860, 90, 1130, 280), "ScheduledPosts", ["schedule_id (PK)", "post_id (FK)", "platform", "scheduled_time", "status"])
    entity((470, 440, 760, 630), "ConnectedAccounts", ["account_id (PK)", "user_id (FK)", "platform", "access_token", "profile_name"])
    arrow(d, (330, 175), (470, 175))
    arrow(d, (760, 180), (860, 180))
    arrow(d, (205, 260), (470, 535))
    d.text((365, 140), "creates", fill="black", font=font(22))
    d.text((790, 140), "scheduled as", fill="black", font=font(22))
    d.text((300, 410), "connects", fill="black", font=font(22))
    d.text((485, 705), "Entity-Relationship Diagram", fill="black", font=font(30, True))
    p2 = ASSET_DIR / "er_diagram.png"
    img.save(p2)
    return p0, p1, p2


def set_cell_shading(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tcPr.append(shd)


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(10)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def style_doc(doc):
    sec = doc.sections[0]
    sec.top_margin = Inches(0.8)
    sec.bottom_margin = Inches(0.75)
    sec.left_margin = Inches(0.9)
    sec.right_margin = Inches(0.9)
    styles = doc.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"].font.size = Pt(11)
    for s in ["Heading 1", "Heading 2", "Heading 3"]:
        styles[s].font.name = "Times New Roman"
        styles[s].font.color.rgb = RGBColor(31, 78, 121)
    styles["Heading 1"].font.size = Pt(16)
    styles["Heading 2"].font.size = Pt(13)
    styles["Heading 3"].font.size = Pt(12)


def para(doc, text="", align=None, bold=False, size=None):
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    r = p.add_run(text)
    r.bold = bold
    if size:
        r.font.size = Pt(size)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.08
    return p


def heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.paragraph_format.space_before = Pt(10 if level > 1 else 16)
    p.paragraph_format.space_after = Pt(6)
    return p


def bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(text)
    p.paragraph_format.space_after = Pt(3)


def table(doc, headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    for i, h in enumerate(headers):
        set_cell_text(t.rows[0].cells[i], h, True)
        set_cell_shading(t.rows[0].cells[i], "D9EAF7")
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            set_cell_text(cells[i], str(val))
    doc.add_paragraph()
    return t


def add_image(doc, path, caption):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Inches(6.7))
    c = doc.add_paragraph(caption)
    c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    c.runs[0].italic = True


def new_page(doc):
    doc.add_page_break()


def build():
    dfd, arch, er = create_diagrams()
    doc = Document()
    style_doc(doc)

    # Cover
    para(doc, "School of Computational Science", WD_ALIGN_PARAGRAPH.CENTER, True, 16)
    para(doc, "Department of Computer Applications", WD_ALIGN_PARAGRAPH.CENTER, True, 14)
    doc.add_paragraph()
    para(doc, PROJECT, WD_ALIGN_PARAGRAPH.CENTER, True, 18)
    para(doc, "PROJECT REPORT", WD_ALIGN_PARAGRAPH.CENTER, True, 16)
    doc.add_paragraph()
    para(doc, "Submitted by:", WD_ALIGN_PARAGRAPH.CENTER, True, 12)
    para(doc, "[Student Name]", WD_ALIGN_PARAGRAPH.CENTER, True, 13)
    para(doc, "[Registration Number]", WD_ALIGN_PARAGRAPH.CENTER, False, 12)
    doc.add_paragraph()
    para(doc, "Under the guidance of", WD_ALIGN_PARAGRAPH.CENTER, False, 12)
    para(doc, "Mr. Suhel Khan", WD_ALIGN_PARAGRAPH.CENTER, True, 13)
    para(doc, "Project Mentor", WD_ALIGN_PARAGRAPH.CENTER, False, 12)
    doc.add_paragraph()
    para(doc, "Presidency University, Bengaluru", WD_ALIGN_PARAGRAPH.CENTER, True, 13)
    para(doc, "May 2026", WD_ALIGN_PARAGRAPH.CENTER, False, 12)
    new_page(doc)

    heading(doc, "BONAFIDE CERTIFICATE", 1)
    para(doc, f"Certified that this project titled \"{PROJECT}\" is a bonafide work carried out by [Student Name], [Registration Number], in partial fulfillment of the requirements for the award of the Master of Computer Applications degree at Presidency University, Bengaluru.", size=11)
    para(doc, "This project has been completed under the guidance and supervision of the project mentor and has not been submitted elsewhere for the award of any degree or diploma.", size=11)
    doc.add_paragraph("\n\n")
    table(doc, ["Project Guide", "Head of Department", "Dean / Associate Dean"], [["Mr. Suhel Khan", "[HOD Name]", "[Dean Name]"]])
    new_page(doc)

    heading(doc, "DECLARATION", 1)
    para(doc, f"I, [Student Name], student of Master of Computer Applications at Presidency University, Bengaluru, hereby declare that the major project titled \"{PROJECT}\" is my original work carried out under the guidance of Mr. Suhel Khan.", size=11)
    para(doc, "The project report has been prepared for academic submission and all sources of information used in the report have been acknowledged appropriately.", size=11)
    doc.add_paragraph("\n")
    para(doc, "PLACE: BENGALURU", bold=True)
    para(doc, "DATE: [Date]", bold=True)
    para(doc, "[Student Name]", WD_ALIGN_PARAGRAPH.RIGHT, True)
    new_page(doc)

    heading(doc, "ACKNOWLEDGMENT", 1)
    ack = [
        "We sincerely thank Presidency University School of Computational Science for giving us the opportunity to create SocialGenie: AI-Powered Social Media Automation Agent.",
        "We would like to express our gratefulness to our project mentor Suhel Khan, who provided valuable guidance, encouragement, and constructive feedback throughout every stage of our project.",
        "We appreciate the support of the Department of Computer Applications for providing the necessary resources, tools, and technical environment required for the successful implementation of this project.",
        "Lastly, we extend our heartfelt gratitude to our families, friends, and fellow classmates, who continuously supported, motivated, and encouraged us during the development of this project.",
    ]
    for x in ack: para(doc, x)
    new_page(doc)

    heading(doc, "ABSTRACT", 1)
    abs_paras = [
        "The purpose of this project was to develop SocialGenie, an AI-powered social media automation agent that helps users generate creative social media content and schedule posts efficiently. In the modern digital environment, businesses, creators, and organizations need to publish regular, engaging, and platform-appropriate content across social media platforms. However, manually creating captions, hashtags, and post ideas can be time-consuming, repetitive, and difficult to manage consistently.",
        "SocialGenie reduces this problem by using artificial intelligence to generate social media captions based on the user's topic or idea. The system allows users to create an account, log in securely, generate AI-based content, view previously generated posts, connect social media accounts, and schedule content for future posting. It combines a React frontend, FastAPI backend, MongoDB database storage, JWT-based authentication, OpenAI-based content generation, and a background worker that monitors scheduled posts.",
        "The project includes user authentication, AI content generation, post history management, scheduling functionality, connected social media account handling, and automated status updates for scheduled posts. LinkedIn posting is implemented as a primary publishing path, while the system structure provides a foundation for future integration with platforms such as Instagram, Facebook, and YouTube.",
        "SocialGenie improves the content creation workflow by reducing manual effort, supporting faster idea generation, and helping users manage social media posts in an organized way. This report presents the purpose, architecture, system analysis, methodology, implementation workflow, database design, testing approach, and future scope of the AI-powered social media automation system.",
    ]
    for x in abs_paras: para(doc, x)
    new_page(doc)

    heading(doc, "TABLE OF CONTENTS", 1)
    toc = [
        ("Chapter 1 - Introduction", "1"),
        ("1.1 Background", "1"), ("1.2 Problem Definition", "2"), ("1.3 Need for the System", "3"), ("1.4 Objectives", "3"), ("1.5 Scope of the Project", "4"), ("1.6 Advantages", "4"), ("1.7 Project Category", "4"), ("1.8 UN SDG Mapping", "5"), ("1.9 Summary", "5"),
        ("Chapter 2 - Literature Review", "6"), ("2.1 Introduction", "6"), ("2.2 Artificial Intelligence in Digital Marketing", "6"), ("2.3 Natural Language Processing and Generative AI", "7"), ("2.4 Social Media Automation", "8"), ("2.5 Content Generation Tools and Existing Systems", "8"), ("2.6 Full-Stack Web Application Technologies", "9"), ("2.7 Comparison with Existing Systems", "10"), ("2.8 Summary", "11"),
        ("Chapter 3 - System Analysis", "12"), ("Chapter 4 - System Design", "17"), ("Chapter 5 - Implementation", "24"), ("Chapter 6 - Testing", "29"), ("Chapter 7 - Results and Discussion", "33"), ("Chapter 8 - Social, Legal, Ethical and Sustainability Aspects", "36"), ("Chapter 9 - Conclusion and Future Scope", "39"), ("References", "41"), ("Appendix", "42"),
    ]
    table(doc, ["Section", "Page"], toc)
    new_page(doc)

    heading(doc, "LIST OF TABLES", 1)
    table(doc, ["Table No.", "Title"], [
        ("3.1", "Functional Requirements"), ("3.2", "Non-Functional Requirements"), ("4.1", "Database Collections"), ("5.1", "Technology Stack"), ("6.1", "Testing Summary"),
    ])
    heading(doc, "LIST OF FIGURES", 1)
    table(doc, ["Figure No.", "Title"], [
        ("4.1", "System Architecture"), ("4.2", "DFD Level-0"), ("4.3", "Entity-Relationship Diagram"),
    ])
    new_page(doc)

    # Chapter 1
    heading(doc, "CHAPTER 1 - INTRODUCTION", 1)
    heading(doc, "1.1 Background", 2)
    for x in [
        "Social media has become an important platform for businesses, creators, and organizations to connect with their audience. Regular posting, attractive captions, relevant hashtags, and timely publishing help improve online visibility and engagement.",
        "However, creating fresh content every day can be time-consuming. Users often spend effort thinking of ideas, writing captions, choosing hashtags, and managing post schedules. Many users also miss posting at the right time, which affects consistency and audience interaction.",
        "Artificial Intelligence can reduce this effort by generating captions and hashtags based on a simple topic entered by the user. The user can then review, edit, publish, or schedule the generated content.",
        f"{PROJECT} was developed to solve this problem by combining AI-based content generation, user authentication, post history, account connection, and scheduling in one web-based system.",
    ]: para(doc, x)
    heading(doc, "1.2 Problem Definition", 2)
    for x in [
        "Even though social media platforms are widely used, creating and managing content regularly is still difficult for many users. The main problem is the effort required to create meaningful captions, choose relevant hashtags, maintain post history, and publish content consistently.",
        "Users often spend time writing captions from scratch. This becomes repetitive when posts need to be created frequently for marketing, branding, or audience engagement. Another issue is scheduling because users may prepare content but forget to publish it at the right time.",
        "The core problem is that no single simple system exists that can help users generate AI-based social media content, manage generated posts, connect social media accounts, and schedule posts through one organized workflow.",
    ]: para(doc, x)
    heading(doc, "1.3 Need for the System", 2)
    for x in [
        "The need for SocialGenie comes from the growing demand for regular and engaging social media content. Users such as creators, small businesses, startups, and organizations often need to create posts frequently, but writing captions and selecting hashtags manually takes time and effort.",
        "A system like SocialGenie provides a simple and organized solution. It allows users to generate AI-based captions, save generated posts, connect social media accounts, and schedule content for future publishing. This reduces manual work and improves productivity.",
    ]: para(doc, x)
    heading(doc, "1.4 Objectives", 2)
    for x in [
        "The objective of SocialGenie is to make social media content creation easier, faster, and more organized. Instead of writing captions and hashtags manually, users can enter a topic and receive AI-generated content that can be reviewed, edited, published, or scheduled.",
        "The system also aims to reduce the time and effort required for regular social media posting, provide secure user login, maintain post history, and support connected account handling.",
    ]: para(doc, x)
    heading(doc, "1.5 Scope of the Project", 2)
    para(doc, "The scope of SocialGenie is focused on AI-based social media content generation and post scheduling. The system allows users to generate captions and hashtags, store generated posts, view previous content, connect social accounts, and schedule posts for future publishing. LinkedIn posting is supported as a primary publishing path, while platforms such as Instagram, Facebook, and YouTube can be expanded in future versions.")
    heading(doc, "1.6 Advantages", 2)
    for x in ["Time saving through quick caption generation.", "Improved consistency through post scheduling.", "Better organization through post history and scheduled content.", "Support for creativity by reducing blank-page effort.", "User-friendly workflow for creators and small teams."]: bullet(doc, x)
    heading(doc, "1.7 Project Category", 2)
    para(doc, "This project belongs to Artificial Intelligence, Natural Language Processing, Full-Stack Web Development, and Digital Marketing Automation. It can be categorized as an AI-powered web-based social media automation system.")
    heading(doc, "1.8 UN SDG Mapping", 2)
    para(doc, "SocialGenie is related to SDG 8: Decent Work and Economic Growth because it improves digital productivity for creators, small businesses, and organizations. It is also related to SDG 9: Industry, Innovation and Infrastructure because it uses AI, web development, automation, and cloud-based storage to build a practical digital tool.")
    heading(doc, "1.9 Summary", 2)
    para(doc, "This chapter introduced the background, problem, need, objectives, scope, advantages, project category, and UN SDG mapping of SocialGenie. The next chapter discusses literature and technologies related to AI-based social media automation.")
    new_page(doc)

    # Chapter 2
    heading(doc, "CHAPTER 2 - LITERATURE REVIEW", 1)
    sections = {
        "2.1 Introduction": "A literature review explains the existing research, technologies, and concepts related to the system. In recent years, Artificial Intelligence has become widely used in digital marketing, content creation, automation, and social media management.",
        "2.2 Artificial Intelligence in Digital Marketing": "Artificial Intelligence helps businesses and creators generate content, improve personalization, understand customer behavior, and manage online campaigns more efficiently. In SocialGenie, AI is used to generate captions and hashtags from user input.",
        "2.3 Natural Language Processing and Generative AI": "Natural Language Processing helps computers understand and generate human language. Generative AI can create new text such as captions, descriptions, replies, and promotional messages. These concepts are central to SocialGenie's content generation module.",
        "2.4 Social Media Automation": "Social media automation uses software tools to manage repetitive tasks such as content planning, post scheduling, publishing, and status tracking. SocialGenie applies automation through scheduling and background monitoring of posts.",
        "2.5 Content Generation Tools and Existing Systems": "Existing tools such as AI writing assistants, Buffer, Hootsuite, Later, and Canva support parts of content creation or scheduling. However, users often need multiple tools for generation, saving, scheduling, and account management.",
        "2.6 Full-Stack Web Application Technologies": "SocialGenie uses React for the frontend, FastAPI for backend APIs, MongoDB for storage, JWT for authentication, and external AI/social media APIs. Together, these technologies support a complete user workflow.",
        "2.7 Comparison with Existing Systems": "Compared with separate AI writing and scheduling tools, SocialGenie combines caption generation, hashtag generation, login, post history, account connection, and scheduling in one focused system.",
        "2.8 Summary": "This chapter reviewed the concepts related to SocialGenie, including AI in marketing, NLP, generative AI, social media automation, existing tools, and full-stack technologies.",
    }
    for h, t in sections.items():
        heading(doc, h, 2)
        para(doc, t)
    new_page(doc)

    # Chapter 3
    heading(doc, "CHAPTER 3 - SYSTEM ANALYSIS", 1)
    for h, body in [
        ("3.1 Introduction", "System analysis helps understand requirements, problems, and the working process of the proposed system. For SocialGenie, the analysis explains the need for an AI-powered tool that can generate content, manage posts, and schedule them efficiently."),
        ("3.2 Existing System", "In the existing approach, users create captions manually or use separate tools for AI writing and scheduling. This increases manual effort and may lead to irregular posting, missed schedules, and poor organization of generated content."),
        ("3.3 Proposed System", "The proposed system provides a single platform where users can register, log in, generate captions and hashtags using AI, view post history, connect social media accounts, and schedule posts for future publishing."),
        ("3.4 Feasibility Study", "SocialGenie is technically feasible because it uses reliable technologies such as React, FastAPI, MongoDB, JWT, and OpenAI. It is operationally feasible because the interface is simple for users. It is economically feasible because most technologies are open-source or available through basic cloud services."),
    ]:
        heading(doc, h, 2)
        para(doc, body)
    heading(doc, "3.5 Functional Requirements", 2)
    table(doc, ["Requirement", "Description"], [
        ("User registration and login", "Users can create accounts and log in securely."),
        ("AI content generation", "The system generates captions and hashtags from user topics."),
        ("Post history", "Generated content is saved and can be viewed later."),
        ("Account connection", "Users can connect supported social media accounts."),
        ("Post scheduling", "Users can schedule generated content for future publishing."),
        ("Status management", "Scheduled posts are tracked and updated by the worker."),
    ])
    heading(doc, "3.6 Non-Functional Requirements", 2)
    table(doc, ["Requirement", "Description"], [
        ("Usability", "The interface should be simple, clear, and responsive."),
        ("Security", "JWT authentication should protect user-specific data."),
        ("Performance", "Captions and data operations should respond in reasonable time."),
        ("Reliability", "Scheduled post status should be tracked properly."),
        ("Maintainability", "The code structure should allow future platform expansion."),
    ])
    heading(doc, "3.7 System Modules", 2)
    for x in ["User Authentication Module", "AI Content Generation Module", "Post History Module", "Scheduling Module", "Social Media Account Connection Module", "Background Worker Module", "Database Module"]:
        bullet(doc, x)
    heading(doc, "3.8 Use Case Model", 2)
    para(doc, "The main actor is the user. The user registers or logs in, enters a topic, generates AI-based content, views post history, connects social media accounts, schedules posts, and checks scheduled post status.")
    heading(doc, "3.9 Summary", 2)
    para(doc, "This chapter explained the existing system, proposed system, feasibility study, requirements, system modules, and use case model. The analysis shows that SocialGenie can reduce manual effort by combining AI caption generation, post history, account connection, and scheduling.")
    new_page(doc)

    # Chapter 4
    heading(doc, "CHAPTER 4 - SYSTEM DESIGN", 1)
    heading(doc, "4.1 System Architecture", 2)
    para(doc, "The architecture of SocialGenie follows a full-stack web application model. The user interacts with the React frontend, which sends API requests to the FastAPI backend. The backend communicates with OpenAI for content generation, MongoDB for storage, social media APIs for account connection and posting, and a background worker for scheduled tasks.")
    add_image(doc, arch, "Figure 4.1: System Architecture")
    heading(doc, "4.2 Data Flow Diagram Level-0", 2)
    para(doc, "The DFD Level-0 shows the system as a single main process. The user sends a topic or scheduling request to SocialGenie. The system communicates with the AI service, database, and social media platform, then returns generated content and status to the user.")
    add_image(doc, dfd, "Figure 4.2: DFD Level-0")
    heading(doc, "4.3 Entity-Relationship Diagram", 2)
    para(doc, "The ER diagram represents the major data entities used in the system. Users create generated posts, generated posts can become scheduled posts, and users can connect social media accounts.")
    add_image(doc, er, "Figure 4.3: Entity-Relationship Diagram")
    heading(doc, "4.4 Database Design / Data Dictionary", 2)
    table(doc, ["Collection", "Purpose", "Key Fields"], [
        ("users", "Stores registered users", "username, password_hash"),
        ("generated_posts", "Stores AI-generated captions", "user_id, topic, caption, hashtags, created_at"),
        ("scheduled_posts", "Stores scheduled content", "user_id, post_id, platform, scheduled_time, status"),
        ("connected_accounts", "Stores account connection details", "user_id, platform, access_token, profile"),
        ("logs", "Stores system activity and errors", "event, user_id, timestamp, status"),
    ])
    heading(doc, "4.5 Flowchart", 2)
    para(doc, "The workflow begins when a user logs in and enters a topic. The backend sends the topic to the AI service, receives generated content, saves it in the database, and displays it to the user. The user may then schedule or publish the post. The background worker checks scheduled posts and updates their status.")
    heading(doc, "4.6 Module Interaction Design", 2)
    para(doc, "The authentication module verifies user access. The content generation module communicates with the AI service. The database module stores user and post records. The scheduling module creates scheduled tasks, and the worker module monitors due posts. Social media service modules handle connected account and posting operations.")
    heading(doc, "4.7 Summary", 2)
    para(doc, "This chapter presented the architecture, DFD, ER diagram, database design, flowchart explanation, and module interactions of SocialGenie.")
    new_page(doc)

    # Chapter 5
    heading(doc, "CHAPTER 5 - IMPLEMENTATION", 1)
    heading(doc, "5.1 Technology Stack", 2)
    table(doc, ["Layer", "Technology", "Purpose"], [
        ("Frontend", "React", "User interface and component-based views"),
        ("Backend", "FastAPI", "REST APIs and business logic"),
        ("Database", "MongoDB", "User, post, schedule, and account storage"),
        ("Authentication", "JWT + bcrypt", "Secure login and password protection"),
        ("AI Service", "OpenAI API", "Caption and hashtag generation"),
        ("Worker", "Python background worker", "Scheduled post monitoring"),
    ])
    heading(doc, "5.2 Frontend Implementation", 2)
    para(doc, "The frontend provides screens for registration, login, dashboard, content generation, post history, account connection, and scheduling. React components make the interface dynamic and easy to update.")
    heading(doc, "5.3 Backend Implementation", 2)
    para(doc, "The FastAPI backend handles routes for authentication, content generation, post retrieval, scheduling, and social media account connection. It validates requests, communicates with external services, and stores results in MongoDB.")
    heading(doc, "5.4 AI Content Generation", 2)
    para(doc, "When the user submits a topic, the backend sends a structured request to the OpenAI service. The generated caption and hashtags are returned to the frontend and saved for future reference.")
    heading(doc, "5.5 Scheduling and Worker Process", 2)
    para(doc, "Scheduled posts are stored with date, time, platform, and status. The worker checks pending schedules at regular intervals and updates records when the scheduled time is reached.")
    heading(doc, "5.6 Summary", 2)
    para(doc, "This chapter explained the implementation of the frontend, backend, AI generation, database storage, authentication, and scheduling workflow.")
    new_page(doc)

    # Chapter 6
    heading(doc, "CHAPTER 6 - TESTING", 1)
    para(doc, "Testing ensures that the system works correctly, securely, and reliably. SocialGenie was tested through functional testing, authentication testing, API testing, database testing, and scheduling workflow testing.")
    table(doc, ["Test Case", "Expected Result", "Status"], [
        ("User registration", "New user account is created", "Pass"),
        ("User login", "Valid user receives access token", "Pass"),
        ("Invalid login", "System rejects wrong credentials", "Pass"),
        ("Generate caption", "Caption and hashtags are generated", "Pass"),
        ("View history", "User sees only own generated posts", "Pass"),
        ("Schedule post", "Post is saved with pending status", "Pass"),
        ("Worker status update", "Due scheduled post status is updated", "Pass"),
        ("Connected account flow", "Platform connection details are handled", "Pass"),
    ])
    heading(doc, "6.1 Summary", 2)
    para(doc, "The testing process shows that the main features of SocialGenie work as expected. Further testing can be added for large-scale users, rate limits, platform API failures, and real publishing workflows.")
    new_page(doc)

    # Chapter 7
    heading(doc, "CHAPTER 7 - RESULTS AND DISCUSSION", 1)
    para(doc, "The project successfully demonstrates an AI-assisted workflow for social media content generation and scheduling. Users can log in, enter a topic, generate captions, store results, connect accounts, and schedule posts through a single interface.")
    para(doc, "The AI content generation module reduces the time required to write captions manually. The scheduling module helps users maintain consistency, while post history provides better organization of previous content. The modular backend also makes the system easier to extend with more social platforms.")
    para(doc, "The result confirms that a lightweight full-stack application can combine AI generation, authentication, database storage, social account handling, and background scheduling in a practical way.")
    new_page(doc)

    # Chapter 8
    heading(doc, "CHAPTER 8 - SOCIAL, LEGAL, ETHICAL AND SUSTAINABILITY ASPECTS", 1)
    heading(doc, "8.1 Social Impact", 2)
    para(doc, "SocialGenie can help small businesses, creators, and students manage digital communication more efficiently by reducing repetitive writing effort.")
    heading(doc, "8.2 Legal Considerations", 2)
    para(doc, "The system should respect platform API terms, copyright rules, privacy requirements, and user consent while connecting social media accounts and publishing content.")
    heading(doc, "8.3 Ethical Considerations", 2)
    para(doc, "AI-generated content should be reviewed by users before publishing. The system should avoid misleading, harmful, or inappropriate content and should not fully replace human judgment.")
    heading(doc, "8.4 Sustainability", 2)
    para(doc, "The project supports efficient digital workflows by reducing repetitive manual work. Future versions should optimize API usage and avoid unnecessary AI calls to reduce computational cost.")
    new_page(doc)

    # Chapter 9
    heading(doc, "CHAPTER 9 - CONCLUSION AND FUTURE SCOPE", 1)
    heading(doc, "9.1 Conclusion", 2)
    para(doc, "SocialGenie presents a practical AI-powered system for social media content generation and scheduling. It combines React, FastAPI, MongoDB, JWT authentication, OpenAI-based content generation, social account connection, and a background worker into one organized workflow.")
    para(doc, "The project reduces manual effort in caption writing, supports post organization, and helps users maintain consistent posting schedules. Its modular design provides a clear foundation for future expansion.")
    heading(doc, "9.2 Future Scope", 2)
    for x in ["Full publishing support for Instagram, Facebook, YouTube, and other platforms.", "Image generation and media upload workflow.", "Analytics dashboard for engagement tracking.", "Brand tone customization and campaign templates.", "Content safety checks and factuality warnings.", "Queue-based scheduler with improved timezone handling."]:
        bullet(doc, x)
    new_page(doc)

    heading(doc, "REFERENCES", 1)
    refs = [
        "OpenAI, OpenAI API Documentation, 2026.",
        "FastAPI, FastAPI Framework Documentation, 2026.",
        "MongoDB, MongoDB Developer Documentation, 2026.",
        "React, React Documentation, 2026.",
        "D. Grewal, C. B. Satornino, T. Davenport, and A. Guha, 'How generative AI is shaping the future of marketing,' Journal of the Academy of Marketing Science, 2025.",
        "P. Cillo and G. Rubera, 'Generative AI in innovation and marketing processes: A roadmap of research opportunities,' Journal of the Academy of Marketing Science, 2025.",
    ]
    for i, r in enumerate(refs, 1):
        para(doc, f"[{i}] {r}")
    new_page(doc)

    heading(doc, "APPENDIX", 1)
    heading(doc, "Appendix A: Sample API Endpoints", 2)
    table(doc, ["Endpoint", "Purpose"], [
        ("/auth/register", "Register a new user"),
        ("/auth/login", "Authenticate user and return token"),
        ("/generate", "Generate AI caption and hashtags"),
        ("/posts/history", "Retrieve generated posts"),
        ("/schedule", "Create scheduled post"),
        ("/accounts/connect", "Connect social media account"),
    ])
    heading(doc, "Appendix B: Editable Student Details", 2)
    para(doc, "Replace [Student Name], [Registration Number], [Date], [HOD Name], and [Dean Name] with the final academic details before submission.")

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
